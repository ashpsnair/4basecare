from docx import Document
import pandas as pd

sheet= pd.read_csv('C:/Users/ash24/Desktop/temp.csv')

document = Document()

document.add_heading('Report', 0)
p = document.add_paragraph('Details about the patient maybe ')

document.add_heading("Patient name : " + str(sheet.iloc[0,1]))
document.add_paragraph('Further details in table', style='Intense Quote')

records = (
    (1, str(sheet.iloc[1,1]), str(sheet.iloc[2,1])),
    (2, str(sheet.iloc[1,2]), str(sheet.iloc[2,2])),
    (3, str(sheet.iloc[1,3]), str(sheet.iloc[2,3]))
)

table = document.add_table(rows=1, cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Serial_no'
hdr_cells[1].text = 'Gene_Id'
hdr_cells[2].text = 'Function'
for qty, id, desc in records:
    row_cells = table.add_row().cells
    row_cells[0].text = str(qty)
    row_cells[1].text = id
    row_cells[2].text = desc


document.add_page_break()

document.save('C:/Users/ash24/Desktop/demo.docx')
